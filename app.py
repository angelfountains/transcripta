import streamlit as st
import whisper
import tempfile
import os
from pydub import AudioSegment
from math import ceil
from docx import Document

# --- CONFIGURACIÓN ---
CHUNK_DURATION = 10 * 60 * 1000  # 10 minutos en milisegundos

# --- TÍTULO ---
st.title("Transcripta")
st.subheader("Transcribe audios largos en partes. Precisión y estabilidad para grabaciones extensas.")

# --- SUBIR ARCHIVO ---
audio_file = st.file_uploader("Sube tu archivo de audio (MP3, WAV, M4A)", type=["mp3", "wav", "m4a"])

if audio_file is not None:
    with st.spinner("Procesando el archivo..."):

        # Guardar audio subido como archivo temporal
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
            tmp.write(audio_file.read())
            tmp_path = tmp.name

        # Dividir en fragmentos de 10 minutos
        audio = AudioSegment.from_file(tmp_path)
        duration_ms = len(audio)
        num_chunks = ceil(duration_ms / CHUNK_DURATION)

        st.info(f"Audio dividido en {num_chunks} parte(s) de 10 minutos.")

        model = whisper.load_model("base")  # o "small" si querés más rapidez

        full_transcription = ""

        for i in range(num_chunks):
            start = i * CHUNK_DURATION
            end = min((i + 1) * CHUNK_DURATION, duration_ms)
            chunk = audio[start:end]

            # Guardar fragmento temporal
            chunk_path = os.path.join(tempfile.gettempdir(), f"chunk_{i}.mp3")
            chunk.export(chunk_path, format="mp3")

            # Transcribir fragmento
            result = model.transcribe(chunk_path, language="es")
            full_transcription += f"\n\n--- Fragmento {i + 1} ---\n\n"
            full_transcription += result["text"]

        # Mostrar resultado
        st.success("Transcripción completa ✅")
        st.text_area("Texto transcrito:", full_transcription, height=400)

        # Descargar como .txt
        st.download_button("Descargar como .txt", full_transcription, file_name="transcripcion.txt")

        # Descargar como .docx
        doc = Document()
        doc.add_heading("Transcripción completa", level=1)
        doc.add_paragraph(full_transcription)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
            doc.save(tmp_docx.name)
            tmp_docx.seek(0)
            st.download_button("Descargar como .docx", tmp_docx.read(), file_name="transcripcion.docx")
